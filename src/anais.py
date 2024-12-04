# **Instalação Inicial de pacotes python** 

pip install pandas python-docx

# **Anais da Mostra**

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

# Caminhos dos arquivos
file_path = '/data/2023_Anais.xlsx'  # Caminho do arquivo Excel
image_path = '/data/logo.png'  # Caminho da imagem para o cabeçalho

# Função para carregar a planilha com tratamento de exceção
def load_excel(file_path, sheet_name):
    try:
        return pd.read_excel(file_path, sheet_name=sheet_name)
    except Exception as e:
        print(f"Erro ao carregar o arquivo Excel: {e}")
        return None

sheet_name1 = 'BD_Mostra'  # Substitua pelo nome da aba desejada

# Carregar a planilha
df = load_excel(file_path, sheet_name1)
if df is None:
    exit()  # Encerra o script se a planilha não puder ser carregada

# Criar um documento Word
doc = Document()

# Função para definir a fonte, estilo e alinhamento
def set_font(run, font_name='Roboto', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

# Função para adicionar a imagem ao cabeçalho e o nome do evento
def add_header(doc, image_path):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Ajuste o tamanho conforme necessário
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhado à direita

    # Adicionar nome do evento logo abaixo da imagem
    p = header.add_paragraph()
    run = p.add_run("VIII Mostra de Ensino, Pesquisa e Extensão")
    set_font(run, font_size=11, bold=True)  # Roboto 11, Bold
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Adicionar Cabeçalho com Imagem e Nome do Evento uma vez
add_header(doc, image_path)

# Função para adicionar a Categoria, Área e Instituição
def add_category_area_institution(doc, category, area, institution):
    p = doc.add_paragraph()

    # Adiciona "Categoria:"
    run = p.add_run("Categoria: ")
    set_font(run, font_size=10, bold=True)  # Negrito até os dois pontos

    # Adiciona o valor da categoria
    run = p.add_run(category)
    set_font(run, font_size=10)  # Roboto 10, normal

    # Adiciona "Área:"
    p.add_run("\n")
    run = p.add_run("Área: ")
    set_font(run, font_size=10, bold=True)  # Negrito até os dois pontos

    # Adiciona o valor da área
    run = p.add_run(area)
    set_font(run, font_size=10)  # Roboto 10, normal

    # Adiciona "Instituição:"
    p.add_run("\n")
    run = p.add_run("Instituição: ")
    set_font(run, font_size=10, bold=True)  # Negrito até os dois pontos

    # Adiciona o valor da instituição
    run = p.add_run(institution)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar o Título do Resumo
def add_title(doc, title_text):
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    set_font(run, font_size=11, bold=True)  # Roboto 11, Bold
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar os Autores e Orientador
def add_authors(doc, author, coauthors, orientador):
    authors_text = f"{author}, {coauthors}, {orientador}"
    p = doc.add_paragraph()
    run = p.add_run(authors_text)
    set_font(run, font_size=10, italic=True)  # Roboto 10, Itálico
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_summary(doc, summary_html):
    soup = BeautifulSoup(summary_html, 'html.parser')

    # Adiciona "Resumo:" no início
    p = doc.add_paragraph()
    run = p.add_run("Resumo: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Concatena todo o texto do HTML em uma única string
    summary_text = ' '.join(paragraph.get_text().strip() for paragraph in soup.find_all('p'))

    # Adiciona o texto concatenado como um único parágrafo, já justificado e formatado
    run = p.add_run(summary_text)
    set_font(run, font_size=10)  # Roboto 10, normal
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificado
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para formatar os ODS
def format_ods(ods):
    # Remove os colchetes e aspas da string e separa por ponto e vírgula
    if pd.isna(ods):
        return ""  # Retorna uma string vazia se o valor for NaN
    ods_list = eval(ods)  # Converte a string para lista
    return "; ".join(ods_list)  # Junta os elementos da lista com ponto e vírgula

# Função para adicionar Palavras-chave e ODS
def add_keywords_ods(doc, keywords, ods):
    p = doc.add_paragraph()

    # Adiciona "Palavras-chave:"
    run = p.add_run("Palavras-chave: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Adiciona o valor das palavras-chave
    run = p.add_run(keywords)
    set_font(run, font_size=10)  # Roboto 10, normal

    # Adiciona "Objetivo de Desenvolvimento Sustentável:"
    p.add_run("\n")
    run = p.add_run("Objetivo de Desenvolvimento Sustentável: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Formata os ODS e adiciona ao documento
    formatted_ods = format_ods(ods)
    run = p.add_run(formatted_ods)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Iterar por cada linha da planilha e gerar uma página por linha
for index, row in df.iterrows():
    # Adicionar Categoria, Área e Instituição
    add_category_area_institution(doc, row['Categoria'], row['Área'], row['Instituição'])

    # Adicionar Título
    add_title(doc, row['Título'])

    # Adicionar Autores e Orientador
    add_authors(doc, row['Autor'], row['Coautores'], row['Orientador'])

    # Adicionar Resumo (formatado com HTML)
    add_summary(doc, row['Resumo'])

    # Adicionar Palavras-Chave e ODS
    add_keywords_ods(doc, row['Palavras-chave'], row['Objetivo de Desenvolvimento Sustentável'])

    # Quebra de página para a próxima linha da planilha
    doc.add_page_break()

# Salvar o documento gerado
try:
    doc.save('/outputs/anais_formatado_mostra.docx')
except Exception as e:
    print(f"Erro ao salvar o documento: {e}")

# **Anais da Congresso**

sheet_name2 = 'BD_Poster'  # Substitua pelo nome da aba desejada

# Carregar a planilha
df = load_excel(file_path, sheet_name2)
if df is None:
    exit()  # Encerra o script se a planilha não puder ser carregada

# Criar um documento Word
doc = Document()

# Função para definir a fonte, estilo e alinhamento
def set_font(run, font_name='Roboto', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

# Função para adicionar a imagem ao cabeçalho e o nome do evento
def add_header(doc, image_path):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Ajuste o tamanho conforme necessário
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhado à direita

# Adicionar Cabeçalho com Imagem e Nome do Evento uma vez
add_header(doc, image_path)

# Função para adicionar a Categoria, Área e Instituição
def add_category_area_institution(doc, institution):
    p = doc.add_paragraph()

    # Adiciona "Instituição:"
    run = p.add_run("Instituição: ")
    set_font(run, font_size=10, bold=True)  # Negrito até os dois pontos

    # Adiciona o valor da instituição
    run = p.add_run(institution)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar o Título do Resumo
def add_title(doc, title_text):
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    set_font(run, font_size=11, bold=True)  # Roboto 11, Bold
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar os Autores e Orientador
def add_authors(doc, author, coauthors, orientador):
    authors_text = f"{author}, {coauthors}, {orientador}"
    p = doc.add_paragraph()
    run = p.add_run(authors_text)
    set_font(run, font_size=10, italic=True)  # Roboto 10, Itálico
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_summary(doc, summary_html):
    soup = BeautifulSoup(summary_html, 'html.parser')

    # Adiciona "Resumo:" no início
    p = doc.add_paragraph()
    run = p.add_run("Resumo: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Concatena todo o texto do HTML em uma única string
    summary_text = ' '.join(paragraph.get_text().strip() for paragraph in soup.find_all('p'))

    # Adiciona o texto concatenado como um único parágrafo, já justificado e formatado
    run = p.add_run(summary_text)
    set_font(run, font_size=10)  # Roboto 10, normal
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificado
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para formatar os ODS
def format_ods(ods):
    # Remove os colchetes e aspas da string e separa por ponto e vírgula
    if pd.isna(ods):
        return ""  # Retorna uma string vazia se o valor for NaN
    ods_list = eval(ods)  # Converte a string para lista
    return "; ".join(ods_list)  # Junta os elementos da lista com ponto e vírgula

# Função para adicionar Palavras-chave e ODS
def add_keywords_ods(doc, keywords, ods):
    p = doc.add_paragraph()

    # Adiciona "Palavras-chave:"
    run = p.add_run("Palavras-chave: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Adiciona o valor das palavras-chave
    run = p.add_run(keywords)
    set_font(run, font_size=10)  # Roboto 10, normal

    # Adiciona "Objetivo de Desenvolvimento Sustentável:"
    p.add_run("\n")
    run = p.add_run("Objetivo de Desenvolvimento Sustentável: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Formata os ODS e adiciona ao documento
    formatted_ods = format_ods(ods)
    run = p.add_run(formatted_ods)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Iterar por cada linha da planilha e gerar uma página por linha
for index, row in df.iterrows():
    # Adicionar Categoria, Área e Instituição
    add_category_area_institution(doc, row['Instituição'])

    # Adicionar Título
    add_title(doc, row['Título'])

    # Adicionar Autores e Orientador
    add_authors(doc, row['Autor'], row['Coautores'], row['Orientador'])

    # Adicionar Resumo (formatado com HTML)
    add_summary(doc, row['Resumo'])

    # Adicionar Palavras-Chave e ODS
    add_keywords_ods(doc, row['Palavras Chave'], row['Objetivo de Desenvolvimento Sustentável'])

    # Quebra de página para a próxima linha da planilha
    doc.add_page_break()

# Salvar o documento gerado
try:
    doc.save('/outputs/anais_formatado_congresso.docx')
except Exception as e:
    print(f"Erro ao salvar o documento: {e}")

# **Anais do Congresso Escolar**

sheet_name3 = 'BD_Connect'  # Substitua pelo nome da aba desejada

# Carregar a planilha
df = load_excel(file_path, sheet_name3)
if df is None:
    exit()  # Encerra o script se a planilha não puder ser carregada

# Criar um documento Word
doc = Document()

# Função para definir a fonte, estilo e alinhamento
def set_font(run, font_name='Roboto', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

# Função para adicionar a imagem ao cabeçalho e o nome do evento
def add_header(doc, image_path):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Ajuste o tamanho conforme necessário
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinhado à direita

    # Adicionar nome do evento logo abaixo da imagem
    p = header.add_paragraph()
    run = p.add_run("2º Congresso Escolar Científico da UFCSPA")
    set_font(run, font_size=11, bold=True)  # Roboto 11, Bold
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Adicionar Cabeçalho com Imagem e Nome do Evento uma vez
add_header(doc, image_path)

# Função para adicionar a Categoria, Área e Instituição
def add_category_area_institution(doc, institution):
    p = doc.add_paragraph()

    # Adiciona "Instituição:"
    run = p.add_run("Instituição: ")
    set_font(run, font_size=10, bold=True)  # Negrito até os dois pontos

    # Adiciona o valor da instituição
    run = p.add_run(institution)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar o Título do Resumo
def add_title(doc, title_text):
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    set_font(run, font_size=11, bold=True)  # Roboto 11, Bold
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para adicionar os Autores
def add_authors(doc, author, coauthors):
    authors_text = f"{author}, {coauthors}"
    p = doc.add_paragraph()
    run = p.add_run(authors_text)
    set_font(run, font_size=10, italic=True)  # Roboto 10, Itálico
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_summary(doc, summary_html):
    soup = BeautifulSoup(summary_html, 'html.parser')

    # Adiciona "Resumo:" no início
    p = doc.add_paragraph()
    run = p.add_run("Resumo: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Concatena todo o texto do HTML em uma única string
    summary_text = ' '.join(paragraph.get_text().strip() for paragraph in soup.find_all('p'))

    # Adiciona o texto concatenado como um único parágrafo, já justificado e formatado
    run = p.add_run(summary_text)
    set_font(run, font_size=10)  # Roboto 10, normal
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificado
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Função para formatar os ODS
def format_ods(ods):
    # Remove os colchetes e aspas da string e separa por ponto e vírgula
    if pd.isna(ods):
        return ""  # Retorna uma string vazia se o valor for NaN
    ods_list = eval(ods)  # Converte a string para lista
    return "; ".join(ods_list)  # Junta os elementos da lista com ponto e vírgula

# Função para adicionar Palavras-chave e ODS
def add_keywords_ods(doc, keywords, ods):
    p = doc.add_paragraph()

    # Adiciona "Palavras-chave:"
    run = p.add_run("Palavras-chave: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Adiciona o valor das palavras-chave
    run = p.add_run(keywords)
    set_font(run, font_size=10)  # Roboto 10, normal

    # Adiciona "Objetivo de Desenvolvimento Sustentável:"
    p.add_run("\n")
    run = p.add_run("Objetivo de Desenvolvimento Sustentável: ")
    set_font(run, font_size=10, bold=True)  # Roboto 10, Bold

    # Formata os ODS e adiciona ao documento
    formatted_ods = format_ods(ods)
    run = p.add_run(formatted_ods)
    set_font(run, font_size=10)  # Roboto 10, normal

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# Iterar por cada linha da planilha e gerar uma página por linha
for index, row in df.iterrows():
    # Adicionar Categoria, Área e Instituição
    add_category_area_institution(doc, row['Instituição'])

    # Adicionar Título
    add_title(doc, row['Título'])

    # Adicionar Autores e Orientador
    add_authors(doc, row['Autor'], row['Coautores'])

    # Adicionar Resumo (formatado com HTML)
    add_summary(doc, row['Resumo'])

    # Adicionar Palavras-Chave e ODS
    add_keywords_ods(doc, row['Palavras Chave'], row['ODS'])

    # Quebra de página para a próxima linha da planilha
    doc.add_page_break()

# Salvar o documento gerado
try:
    doc.save('/outputs/anais_formatado_escolar.docx')
except Exception as e:
    print(f"Erro ao salvar o documento: {e}")